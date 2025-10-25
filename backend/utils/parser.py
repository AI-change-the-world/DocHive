from typing import Optional
import io
from pathlib import Path
import PyPDF2
import docx
from PIL import Image
from config import get_settings

settings = get_settings()


class DocumentParser:
    """文档解析器"""

    @staticmethod
    async def parse_pdf(file_data: bytes) -> str:
        """解析 PDF 文件"""
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"PDF 解析失败: {str(e)}")

    @staticmethod
    async def parse_docx(file_data: bytes) -> str:
        """解析 DOCX 文件"""
        try:
            doc_file = io.BytesIO(file_data)
            doc = docx.Document(doc_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        text_content.append(row_text)

            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"DOCX 解析失败: {str(e)}")

    @staticmethod
    async def parse_txt(file_data: bytes, encoding: str = "utf-8") -> str:
        """解析文本文件"""
        try:
            # 尝试多种编码
            encodings = [encoding, "utf-8", "gbk", "gb2312", "latin-1"]

            for enc in encodings:
                try:
                    return file_data.decode(enc)
                except UnicodeDecodeError:
                    continue

            # 如果都失败，使用默认编码并忽略错误
            return file_data.decode("utf-8", errors="ignore")
        except Exception as e:
            raise Exception(f"文本文件解析失败: {str(e)}")

    @staticmethod
    async def parse_image_ocr(file_data: bytes) -> str:
        """使用 OCR 解析图片, 后续会增加大模型OCR"""
        raise NotImplementedError("暂不支持图片解析")

    @staticmethod
    async def parse_file(file_data: bytes, file_extension: str) -> str:
        """
        根据文件扩展名解析文件

        Args:
            file_data: 文件二进制数据
            file_extension: 文件扩展名（如 .pdf, .docx）

        Returns:
            解析后的文本内容
        """
        ext = file_extension.lower().lstrip(".")

        if ext == "pdf":
            return await DocumentParser.parse_pdf(file_data)
        elif ext in ["docx", "doc"]:
            return await DocumentParser.parse_docx(file_data)
        elif ext in ["txt", "md", "markdown"]:
            return await DocumentParser.parse_txt(file_data)
        elif ext in ["png", "jpg", "jpeg", "bmp", "tiff"]:
            return await DocumentParser.parse_image_ocr(file_data)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def extract_metadata(file_data: bytes, file_extension: str) -> dict:
        """提取文件元信息"""
        metadata = {
            "file_size": len(file_data),
            "file_type": file_extension,
        }

        try:
            ext = file_extension.lower().lstrip(".")

            if ext == "pdf":
                pdf_file = io.BytesIO(file_data)
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                metadata.update(
                    {
                        "pages": len(pdf_reader.pages),
                        "pdf_metadata": (
                            pdf_reader.metadata if pdf_reader.metadata else {}
                        ),
                    }
                )

            elif ext in ["docx", "doc"]:
                doc_file = io.BytesIO(file_data)
                doc = docx.Document(doc_file)

                core_props = doc.core_properties
                metadata.update(
                    {
                        "author": core_props.author or "",
                        "created": (
                            str(core_props.created) if core_props.created else ""
                        ),
                        "modified": (
                            str(core_props.modified) if core_props.modified else ""
                        ),
                        "title": core_props.title or "",
                    }
                )

        except Exception as e:
            metadata["metadata_error"] = str(e)

        return metadata
